import os
from dotenv import load_dotenv
from pinecone import Pinecone
import docx
from sentence_transformers import SentenceTransformer
# 1. Cargar las variables del archivo .env
load_dotenv()
# 2. Obtener la API Key de forma segura
pinecone_api_key = os.getenv('PINECONE_API_KEY')
# 3. Inicializar el cliente oficial con 'P' mayúscula
pc = Pinecone(api_key=pinecone_api_key)
# 4. Apuntar a tu índice (usando la 'i' minúscula en index)
index = pc.index("my-db-rag")
doc_path = "Manual de Contingencias.docx"
doc = docx.Document(doc_path)
texto_completo = []
for parrafo in doc.paragraphs:
    if parrafo.text.strip():  # Evita guardar líneas vacías
        texto_completo.append(parrafo.text)

print(f"¡Manual leído con éxito! Se encontraron {len(texto_completo)} párrafos de texto.")
model=SentenceTransformer('all-mpnet-base-v2')
doc = docx.Document(doc_path)
vector_to_upload=[]
for i, parrafo in enumerate(doc.paragraphs):
    texto = parrafo.text.strip()

    if texto:  # Si el párrafo no está vacío
        # Generar la lista de 768 números basada en el texto
        embedding = model.encode(texto).tolist()

        # Estructurar el paquete como lo exige Pinecone
        vector_empaquetado = {
            "id": f"regla_{i}",  # Un ID único para cada párrafo
            "values": embedding,  # El vector de 768 números
            "metadata": {"text": texto}  # El texto original en español para que la LLM lo lea después
        }
        vector_to_upload.append(vector_empaquetado)

# 8. Hacer el subido (Upsert) oficial a Pinecone
print(f"Subiendo {len(vector_to_upload)} fragmentos a Pinecone...")
index.upsert(vectors=vector_to_upload)
print("¡Funciona!")
